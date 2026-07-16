"""Runnable self-check for NeuRAG embedder + search ranking. Stdlib only.

    python selfcheck.py         (run from inside the NeuRAG/ folder — no install needed)

Forces the lexical path (NEURAG_EMBEDDER=null) so it is deterministic and needs
no model download. Raises AssertionError on regression, else prints ALL OK.
"""

import os
import sys
import tempfile
from pathlib import Path

os.environ["NEURAG_EMBEDDER"] = "null"  # force lexical, deterministic, offline
# Make `neurag` importable when run directly, without pip install -e.
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from neurag.embedder import get_embedder, NullEmbedder, DIM
from neurag.db import KnowledgeGraph
from neurag.chunker import chunk_file


def check_embedder() -> None:
    e = get_embedder()
    assert e.name == "null" and e.embed("x") is None
    assert NullEmbedder().dim == DIM
    print("OK embedder: null routing + None embed")


def check_lexical_search() -> None:
    tmp = Path(tempfile.mkdtemp()) / "kg.db"
    kg = KnowledgeGraph(tmp)
    god = kg.add_node("Java", "godnode")
    n = kg.add_node("Concurrency", "fundamental", parent_id=god)
    kg.add_chunk(n, "Threads and locks manage concurrent access in the JVM.", section="a", chunk_index=0)
    kg.add_chunk(n, "Garbage collection reclaims unused heap memory.", section="b", chunk_index=1)
    kg.add_chunk(n, "A ForkJoinPool schedules parallel tasks across threads.", section="c", chunk_index=2)

    top = kg.search("threads concurrent", top_n=2)
    assert top, "search returned nothing"
    # chunk 'a' matches both query terms → must rank first (over 'c' which matches one)
    assert "concurrent" in top[0]["text"].lower(), f"wrong top hit: {top[0]['text']!r}"
    # the off-topic GC chunk must not surface in top-2
    assert all("garbage" not in c["text"].lower() for c in top), "off-topic chunk ranked"
    print("OK search: TF-IDF ranks the on-topic chunk first")
    kg.close()


def check_docx_chunker() -> None:
    try:
        import docx  # python-docx
    except ImportError:
        print("SKIP docx: python-docx not installed (pip install neurag[docx])")
        return
    d = docx.Document()
    d.add_heading("Setup", level=1)
    d.add_paragraph("Install the runtime and configure the token.")
    d.add_heading("Usage", level=1)
    d.add_paragraph("Call the endpoint with your query to retrieve results.")
    tmp = Path(tempfile.mkdtemp()) / "doc.docx"
    d.save(str(tmp))
    sections = [c.section for c in chunk_file(tmp)]
    assert "Setup" in sections and "Usage" in sections, f"docx sections wrong: {sections}"
    print("OK docx: heading-based sections")


def check_yaml_import() -> None:
    try:
        import yaml  # PyYAML
    except ImportError:
        print("SKIP yaml import: PyYAML not installed (pip install neurag[yaml])")
        return
    from neurag.importer import import_mapping
    d = Path(tempfile.mkdtemp())
    (d / "threads.md").write_text("## Threads\nThreads and locks manage concurrent access.\n", encoding="utf-8")
    (d / "map.yaml").write_text(
        "godnode: Java\n"
        "nodes:\n"
        "  - name: Concurrency\n"
        "    type: fundamental\n"
        "    parent: Java\n"
        "    triggers: [thread, lock]\n"
        "    sources: [threads.md]\n",
        encoding="utf-8",
    )
    kg = KnowledgeGraph(d / "kg.db")
    report = import_mapping(kg, d / "map.yaml")
    assert report["nodes"] == 2, report          # Java godnode + Concurrency
    assert report["chunks"] >= 1, report
    assert kg.get_node_by_name("Concurrency"), "node not created"
    kg.close()
    print("OK yaml import: nodes + chunks from mapping")


if __name__ == "__main__":
    check_embedder()
    check_lexical_search()
    check_docx_chunker()
    check_yaml_import()
    print("ALL OK")
